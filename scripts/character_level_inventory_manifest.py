#!/usr/bin/env python3
"""Build a lightweight character-level inventory manifest for Jyotish sources.

This is an index and governance manifest, not a semantic promotion pass. It
hashes every in-scope file, extracts text metadata when cheap, and queues heavy
PDF/image/document extraction instead of running expensive OCR by default.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from xml.etree import ElementTree
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
REPORT_JSON = ROOT / "docs" / "research" / "character_level_inventory_manifest_latest.json"
REPORT_MD = ROOT / "docs" / "research" / "character_level_inventory_manifest_latest.md"
EXTERNAL_REPORT_JSON = ROOT / "docs" / "research" / "character_level_external_manifest_latest.json"
EXTERNAL_REPORT_MD = ROOT / "docs" / "research" / "character_level_external_manifest_latest.md"
EXTRACTION_QUEUE_JSON = ROOT / "docs" / "research" / "character_level_extraction_queue_latest.json"
EXTRACTION_QUEUE_MD = ROOT / "docs" / "research" / "character_level_extraction_queue_latest.md"
EXTRACTION_RESULTS_JSON = ROOT / "docs" / "research" / "character_level_extraction_results_latest.json"
EXTRACTION_RESULTS_MD = ROOT / "docs" / "research" / "character_level_extraction_results_latest.md"
VISION_OCR_SOURCE = r'''
import Foundation
import Vision
import AppKit

let args = CommandLine.arguments
if args.count < 2 {
    fputs("usage: vision_ocr image\n", stderr)
    exit(2)
}
let url = URL(fileURLWithPath: args[1])
guard let image = NSImage(contentsOf: url), let tiff = image.tiffRepresentation,
      let bitmap = NSBitmapImageRep(data: tiff), let cgImage = bitmap.cgImage else {
    fputs("cannot load image\n", stderr)
    exit(3)
}
let request = VNRecognizeTextRequest { request, error in
    if let error = error {
        fputs("vision error: \(error)\n", stderr)
        exit(4)
    }
    let observations = request.results as? [VNRecognizedTextObservation] ?? []
    for obs in observations {
        if let top = obs.topCandidates(1).first {
            print(top.string)
        }
    }
}
request.recognitionLevel = .accurate
request.usesLanguageCorrection = true
if #available(macOS 11.0, *) {
    request.recognitionLanguages = ["zh-Hans", "en-US"]
}
let handler = VNImageRequestHandler(cgImage: cgImage, options: [:])
do {
    try handler.perform([request])
} catch {
    fputs("perform error: \(error)\n", stderr)
    exit(5)
}
'''

PROJECT_SCAN_ROOTS = [
    "references",
    "references/open_source_sources",
    "docs/research",
    "SKILL.md",
    "AGENTS.md",
]

EXTERNAL_SCAN_ROOTS = [
    Path.home() / "Downloads",
    Path.home() / "Desktop",
    Path.home() / "文件仓库" / "印度占星文章",
    Path.home() / ".workbuddy" / "skills" / "jyotish-vedic-astrology",
    Path.home() / "WorkBuddy" / "2026-06-09-20-03-34" / "jyotish-fragments",
    Path.home() / "WorkBuddy" / "2026-06-10-21-30-47",
    Path.home() / "WorkBuddy" / "2026-06-12-15-22-12",
    Path.home() / "WorkBuddy" / "20260422235041",
    Path.home() / "WorkBuddy" / "20260503121822",
    Path.home() / "WorkBuddy" / "engines-repo" / "jyotish",
    Path.home() / "engines-repo" / "jyotish",
    Path.home() / "Documents" / "ObsidianVault" / "03_研究_术数占星",
]

EXTERNAL_KEYWORDS = [
    "印度占星",
    "占星",
    "jyotish",
    "vedic",
    "astrology",
    "jhora",
    "pyjhora",
    "vedastro",
    "dasha",
    "shadbala",
    "tajika",
    "jaimini",
]

EXCLUDED_PARTS = {
    ".git",
    ".pytest_cache",
    ".ruff_cache",
    "__pycache__",
    "node_modules",
    "venv",
    "venv_vedastro",
    ".venv",
    "build",
    "dist",
}

TEXT_EXTENSIONS = {
    ".css",
    ".csv",
    ".html",
    ".ini",
    ".js",
    ".json",
    ".jsonl",
    ".jsx",
    ".mjs",
    ".md",
    ".py",
    ".sh",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".yaml",
    ".yml",
}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
DOCUMENT_EXTENSIONS = {".doc", ".docx", ".epub", ".mobi", ".pages", ".ppt", ".pptx", ".rtf", ".xls", ".xlsx"}
QUEUED_EXTRACTION_STATUSES = {
    "pdf_text_extraction_queued",
    "image_ocr_queued",
    "document_text_extraction_queued",
}


def _relative(path: Path) -> str:
    return str(path.relative_to(ROOT))


def _should_skip(path: Path) -> bool:
    try:
        rel_parts = path.relative_to(ROOT).parts
    except ValueError:
        rel_parts = path.parts
    return any(part in EXCLUDED_PARTS for part in rel_parts)


def _iter_project_files() -> list[Path]:
    files: set[Path] = set()
    for root_name in PROJECT_SCAN_ROOTS:
        root = ROOT / root_name
        if not root.exists() or _should_skip(root):
            continue
        if root.is_file():
            files.add(root)
        else:
            files.update(path for path in root.rglob("*") if path.is_file() and not _should_skip(path))
    return sorted(files)


def _is_external_relevant(path: Path) -> bool:
    lowered = str(path).lower()
    return any(keyword.lower() in lowered for keyword in EXTERNAL_KEYWORDS)


def _iter_external_files() -> list[Path]:
    files: set[Path] = set()
    for root in EXTERNAL_SCAN_ROOTS:
        if not root.exists() or _should_skip(root):
            continue
        if root.is_file():
            if _is_external_relevant(root):
                files.add(root)
            continue
        for path in root.rglob("*"):
            if path.is_file() and not _should_skip(path) and _is_external_relevant(path):
                files.add(path)
    return sorted(files)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _text_stats(path: Path) -> dict[str, Any]:
    raw = path.read_bytes()
    decoded = raw.decode("utf-8", errors="replace")
    replacement_count = decoded.count("\ufffd")
    return {
        "extraction_status": "text_decode_lossy" if replacement_count else "text_indexed",
        "character_count": len(decoded),
        "line_count": decoded.count("\n") + (1 if decoded else 0),
        "decode_replacement_count": replacement_count,
    }


def _extraction_status(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _text_stats(path)
    if suffix in PDF_EXTENSIONS:
        return {
            "extraction_status": "pdf_text_extraction_queued",
            "character_count": None,
            "line_count": None,
            "decode_replacement_count": None,
        }
    if suffix in IMAGE_EXTENSIONS:
        return {
            "extraction_status": "image_ocr_queued",
            "character_count": None,
            "line_count": None,
            "decode_replacement_count": None,
        }
    if suffix in DOCUMENT_EXTENSIONS:
        return {
            "extraction_status": "document_text_extraction_queued",
            "character_count": None,
            "line_count": None,
            "decode_replacement_count": None,
        }
    return {
        "extraction_status": "binary_indexed",
        "character_count": None,
        "line_count": None,
        "decode_replacement_count": None,
    }


def _classify_path(path: str) -> dict[str, str]:
    if path.startswith("references/open_source_sources/rishi-ai-mcp/"):
        return {
            "classification": "open_source_reference",
            "priority": "priority_1",
            "promotion_status": "reference_layer_candidate",
            "reason": "User-prioritized open-source workflow corpus; requires license-aware review.",
        }
    if path.startswith("references/open_source_sources/vedic-astro-skills/"):
        return {
            "classification": "open_source_reference",
            "priority": "priority_1",
            "promotion_status": "reference_layer_candidate",
            "reason": "User-prioritized Jyotish skill corpus; requires selective source promotion.",
        }
    if path.startswith("references/open_source_sources/"):
        return {
            "classification": "open_source_reference",
            "priority": "priority_2",
            "promotion_status": "reference_layer_candidate",
            "reason": "Open-source corpus; index first, promote only after license and conflict review.",
        }
    if path.startswith("references/real_case_studies/"):
        return {
            "classification": "real_case_calibration",
            "priority": "priority_1",
            "promotion_status": "reference_layer_candidate",
            "reason": "Real-case material for calibration; not direct rule truth.",
        }
    if path.startswith("references/oracle/"):
        return {
            "classification": "oracle_artifact",
            "priority": "priority_2",
            "promotion_status": "oracle_evidence_only",
            "reason": "Oracle artifact; use for parity evidence, not interpretive rules.",
        }
    if path.startswith("references/"):
        return {
            "classification": "reference_candidate",
            "priority": "priority_1",
            "promotion_status": "reference_layer_candidate",
            "reason": "Local reference source candidate; requires content-level review before promotion.",
        }
    if path.startswith("docs/research/local_drafts/"):
        return {
            "classification": "quarantined_draft",
            "priority": "priority_3",
            "promotion_status": "not_truth_source",
            "reason": "Historical draft; do not promote without explicit review.",
        }
    if path.startswith("docs/research/"):
        return {
            "classification": "research_governance",
            "priority": "priority_3",
            "promotion_status": "governance_or_history",
            "reason": "Research history or governance record; index separately from rule truth.",
        }
    return {
        "classification": "project_governance",
        "priority": "priority_3",
        "promotion_status": "governance_or_history",
        "reason": "Project-level governance file.",
    }


def _root_bucket(path: str) -> str:
    if path.startswith("references/open_source_sources/"):
        return "references/open_source_sources"
    if path.startswith("references/"):
        return "references"
    if path.startswith("docs/research/"):
        return "docs/research"
    return path.split("/", 1)[0]


def _external_path(path: Path) -> str:
    return str(path)


def _classify_external_path(path: str) -> dict[str, str]:
    suffix = Path(path).suffix.lower()
    if "/.workbuddy/skills/jyotish-vedic-astrology/" in path:
        return {
            "classification": "external_skill_fragment",
            "priority": "priority_1",
            "promotion_status": "external_candidate",
            "reason": "External Jyotish skill copy; index only until diffed against repo truth.",
        }
    if "/engines-repo/jyotish/" in path or "/WorkBuddy/engines-repo/jyotish/" in path:
        return {
            "classification": "external_engine_fragment",
            "priority": "priority_1",
            "promotion_status": "external_candidate",
            "reason": "External engine fragment; requires source diff and license boundary review.",
        }
    if "/WorkBuddy/" in path or "/Desktop/" in path:
        return {
            "classification": "external_historical_report",
            "priority": "priority_2",
            "promotion_status": "external_reference_only",
            "reason": "Historical report or work artifact; do not promote without privacy review.",
        }
    if suffix in PDF_EXTENSIONS or suffix in DOCUMENT_EXTENSIONS or "/文件仓库/印度占星文章/" in path:
        return {
            "classification": "external_book_or_document",
            "priority": "priority_1",
            "promotion_status": "external_reference_only",
            "reason": "External book/document source; requires extraction and source grading before use.",
        }
    return {
        "classification": "external_archive_or_binary",
        "priority": "priority_3",
        "promotion_status": "external_index_only",
        "reason": "External relevant asset; indexed without content promotion.",
    }


def _external_root_bucket(path: str) -> str:
    home = str(Path.home())
    if path.startswith(home + "/Downloads/"):
        return "~/Downloads"
    if path.startswith(home + "/Desktop/"):
        return "~/Desktop"
    if path.startswith(home + "/文件仓库/"):
        return "~/文件仓库"
    if path.startswith(home + "/.workbuddy/"):
        return "~/.workbuddy"
    if path.startswith(home + "/WorkBuddy/"):
        return "~/WorkBuddy"
    if path.startswith(home + "/engines-repo/"):
        return "~/engines-repo"
    if path.startswith(home + "/Documents/ObsidianVault/"):
        return "~/Documents/ObsidianVault"
    return path


def _build_index(
    *,
    files: list[Path],
    path_label,
    classify,
    root_bucket,
) -> tuple[dict[str, dict[str, Any]], dict[str, Any], dict[str, dict[str, int]]]:
    by_path: dict[str, dict[str, Any]] = {}
    classification_counts: dict[str, int] = {}
    extraction_counts: dict[str, int] = {}
    root_summary: dict[str, dict[str, int]] = {}
    unclassified_files = 0
    unknown_extraction_status = 0

    for file_path in files:
        rel = path_label(file_path)
        file_class = classify(rel)
        extraction = _extraction_status(file_path)
        if not file_class.get("classification"):
            unclassified_files += 1
        if not extraction.get("extraction_status"):
            unknown_extraction_status += 1
        byte_count = file_path.stat().st_size
        item = {
            "path": rel,
            "suffix": file_path.suffix.lower(),
            "byte_count": byte_count,
            "sha256": _sha256(file_path),
            **extraction,
            **file_class,
        }
        by_path[rel] = item
        classification_counts[item["classification"]] = classification_counts.get(item["classification"], 0) + 1
        extraction_counts[item["extraction_status"]] = extraction_counts.get(item["extraction_status"], 0) + 1
        bucket = root_bucket(rel)
        root_summary.setdefault(bucket, {"files": 0, "bytes": 0})
        root_summary[bucket]["files"] += 1
        root_summary[bucket]["bytes"] += byte_count

    summary = {
        "total_files": len(by_path),
        "unhashed_files": sum(1 for item in by_path.values() if not item.get("sha256")),
        "unclassified_files": unclassified_files,
        "unknown_extraction_status": unknown_extraction_status,
        "classification_counts": dict(sorted(classification_counts.items())),
        "extraction_status_counts": dict(sorted(extraction_counts.items())),
    }
    return by_path, summary, dict(sorted(root_summary.items()))


def build_manifest(*, scope: str = "project", write: bool = True) -> dict[str, Any]:
    if scope == "project":
        files = _iter_project_files()
        path_label = _relative
        classify = _classify_path
        root_bucket = _root_bucket
        scan_roots = PROJECT_SCAN_ROOTS
        json_report = REPORT_JSON
        markdown_report = REPORT_MD
        title = "Character-Level Inventory Manifest"
        mode = {
            "heavy_ocr": False,
            "whole_machine_scan": False,
            "external_high_relevance_scan": False,
            "semantic_promotion": False,
            "boundary": "Fast manifest only: hashes and cheap text stats now; heavy OCR and whole-machine scan are queued.",
        }
    elif scope == "external":
        files = _iter_external_files()
        path_label = _external_path
        classify = _classify_external_path
        root_bucket = _external_root_bucket
        scan_roots = [str(path) for path in EXTERNAL_SCAN_ROOTS]
        json_report = EXTERNAL_REPORT_JSON
        markdown_report = EXTERNAL_REPORT_MD
        title = "External High-Relevance Inventory Manifest"
        mode = {
            "heavy_ocr": False,
            "whole_machine_scan": False,
            "external_high_relevance_scan": True,
            "semantic_promotion": False,
            "boundary": "External scan is high-relevance only. It records hashes and extraction states without copying private source text into repo truth.",
        }
    elif scope == "extraction-queue":
        return build_extraction_queue(write=write)
    elif scope == "extraction-results":
        return build_extraction_results(write=write)
    else:
        raise ValueError(f"Unsupported scope: {scope}")

    by_path, summary, root_summary = _build_index(
        files=files,
        path_label=path_label,
        classify=classify,
        root_bucket=root_bucket,
    )
    status = (
        "pass"
        if summary["unhashed_files"] == 0
        and summary["unclassified_files"] == 0
        and summary["unknown_extraction_status"] == 0
        else "fail"
    )
    report = {
        "scope": scope,
        "status": status,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "mode": mode,
        "scan_roots": scan_roots,
        "summary": summary,
        "root_summary": root_summary,
        "by_path": by_path,
        "artifacts": {
            "json_report": _relative(json_report),
            "markdown_report": _relative(markdown_report),
        },
        "title": title,
    }
    if write:
        _write_reports(report, json_report=json_report, markdown_report=markdown_report)
    return report


def _hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()


def _extract_docx_text(path: Path) -> tuple[str, str | None]:
    try:
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                table_text.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs + table_text), None
    except Exception as exc:  # pragma: no cover - exercised by real files.
        fallback_text, fallback_error = _extract_docx_xml_text(path)
        if fallback_text.strip():
            return fallback_text, f"python-docx fallback used after {type(exc).__name__}: {exc}"
        return "", fallback_error or f"{type(exc).__name__}: {exc}"


def _extract_docx_xml_text(path: Path) -> tuple[str, str | None]:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_bytes = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml_bytes)
        text_nodes = []
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for node in root.iter(f"{namespace}t"):
            if node.text:
                text_nodes.append(node.text)
        return "\n".join(text_nodes), None
    except Exception as exc:
        return "", f"docx_xml_fallback_failed {type(exc).__name__}: {exc}"


def _extract_pdf_text(path: Path) -> tuple[str, str, str | None]:
    try:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        text = "\n".join(parts)
        if text.strip():
            return text, "pdfplumber", None
    except Exception as exc:
        pdfplumber_error = f"{type(exc).__name__}: {exc}"
    else:
        pdfplumber_error = None
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, "pypdf", None
    except Exception as exc:  # pragma: no cover - depends on local PDF shape.
        errors = "; ".join(error for error in [pdfplumber_error, f"{type(exc).__name__}: {exc}"] if error)
        return "", "pypdf", errors


def _extract_image_text(path: Path) -> tuple[str, str, str | None]:
    if not shutil.which("tesseract"):
        return _extract_image_text_with_vision(path)
    try:
        from PIL import Image
        import pytesseract

        with Image.open(path) as image:
            return pytesseract.image_to_string(image, lang="+".join(_ocr_languages())), "pytesseract", None
    except Exception as exc:  # pragma: no cover - depends on local OCR install/languages.
        return "", "pytesseract", f"{type(exc).__name__}: {exc}"


def _vision_ocr_binary() -> Path | None:
    if sys.platform != "darwin" or not shutil.which("swiftc"):
        return None
    cache_dir = Path.home() / ".cache" / "jyotish-ocr"
    cache_dir.mkdir(parents=True, exist_ok=True)
    source = cache_dir / "vision_ocr.swift"
    binary = cache_dir / "vision_ocr"
    if binary.exists() and binary.stat().st_mtime >= source.stat().st_mtime if source.exists() else False:
        return binary
    source.write_text(VISION_OCR_SOURCE, encoding="utf-8")
    completed = subprocess.run(
        ["swiftc", str(source), "-o", str(binary)],
        text=True,
        capture_output=True,
        timeout=60,
        check=False,
    )
    if completed.returncode != 0:
        return None
    return binary


def _extract_image_text_with_vision(path: Path) -> tuple[str, str, str | None]:
    binary = _vision_ocr_binary()
    if not binary:
        return "", "none", "no OCR backend available: tesseract missing and macOS Vision unavailable"
    completed = subprocess.run(
        [str(binary), str(path)],
        text=True,
        capture_output=True,
        timeout=90,
        check=False,
    )
    if completed.returncode != 0:
        return "", "macos_vision", completed.stderr.strip() or f"vision_ocr exit {completed.returncode}"
    return completed.stdout, "macos_vision", None


def _ocr_cache_path(item: dict[str, Any], method: str) -> Path:
    cache_dir = Path.home() / ".cache" / "jyotish-ocr" / "results"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / f"{item['sha256']}.{method}.json"


def _load_ocr_cached_item(item: dict[str, Any]) -> dict[str, Any] | None:
    for method in ["macos_vision", "pytesseract"]:
        path = _ocr_cache_path(item, method)
        if not path.exists():
            continue
        try:
            cached = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if cached.get("source_sha256") != item.get("sha256"):
            continue
        return {
            **item,
            "extraction_result": cached["extraction_result"],
            "extraction_method": cached["extraction_method"],
            "blocked_reason": cached.get("blocked_reason"),
            "extracted_character_count": cached["extracted_character_count"],
            "extracted_line_count": cached["extracted_line_count"],
            "text_sha256": cached.get("text_sha256"),
            "post_extraction_classification": cached["post_extraction_classification"],
            "ocr_engine_available": cached["ocr_engine_available"],
            "ocr_requested_languages": cached["ocr_requested_languages"],
            "ocr_available_languages": cached["ocr_available_languages"],
            "ocr_backend": cached["ocr_backend"],
            "ocr_cache_status": "hit",
        }
    return None


def _write_ocr_cached_item(item: dict[str, Any]) -> None:
    method = str(item.get("ocr_backend") or item.get("extraction_method") or "unknown")
    path = _ocr_cache_path(item, method)
    cache_payload = {
        "source_sha256": item["sha256"],
        "extraction_result": item["extraction_result"],
        "extraction_method": item["extraction_method"],
        "blocked_reason": item.get("blocked_reason"),
        "extracted_character_count": item["extracted_character_count"],
        "extracted_line_count": item["extracted_line_count"],
        "text_sha256": item.get("text_sha256"),
        "post_extraction_classification": item["post_extraction_classification"],
        "ocr_engine_available": item.get("ocr_engine_available"),
        "ocr_requested_languages": item.get("ocr_requested_languages", ["chi_sim", "eng"]),
        "ocr_available_languages": item.get("ocr_available_languages", []),
        "ocr_backend": item.get("ocr_backend", method),
    }
    path.write_text(json.dumps(cache_payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _ocr_available_languages() -> list[str]:
    if not shutil.which("tesseract"):
        return []
    try:
        import pytesseract

        return sorted(str(lang) for lang in pytesseract.get_languages(config=""))
    except Exception:
        return []


def _ocr_languages() -> list[str]:
    available = set(_ocr_available_languages())
    preferred = ["chi_sim", "eng"]
    selected = [lang for lang in preferred if lang in available]
    return selected or (["eng"] if "eng" in available else preferred)


def _post_extraction_classification(item: dict[str, Any]) -> str:
    path = str(item["path"])
    if "/Desktop/" in path or "/WorkBuddy/" in path:
        return "extracted_private_reference_only"
    if item["classification"] in {"external_book_or_document", "reference_candidate"}:
        return "extracted_candidate_for_review"
    return "extracted_reference_only"


def _extract_queued_item(item: dict[str, Any]) -> dict[str, Any]:
    path = Path(item["path"])
    status = item["extraction_status"]
    if status == "document_text_extraction_queued" and path.suffix.lower() == ".docx":
        text, error = _extract_docx_text(path)
        method = "docx"
    elif status == "pdf_text_extraction_queued":
        text, method, error = _extract_pdf_text(path)
    elif status == "image_ocr_queued":
        cached = _load_ocr_cached_item(item)
        if cached:
            return cached
        text, method, error = _extract_image_text(path)
        if error and "no OCR backend available" in error:
            blocked_item = {
                **item,
                "extraction_result": "ocr_blocked_missing_engine",
                "extraction_method": method,
                "blocked_reason": error,
                "extracted_character_count": 0,
                "extracted_line_count": 0,
                "text_sha256": None,
                "post_extraction_classification": "extracted_reference_only",
                "ocr_engine_available": False,
                "ocr_requested_languages": ["chi_sim", "eng"],
                "ocr_available_languages": [],
                "ocr_backend": "none",
                "ocr_cache_status": "miss",
            }
            _write_ocr_cached_item(blocked_item)
            return blocked_item
    else:
        text = ""
        method = "unsupported"
        error = "unsupported extraction target"

    normalized = text or ""
    if error and not normalized.strip():
        result = "extraction_failed"
    elif normalized.strip():
        result = "text_extracted"
    else:
        result = "text_empty"
    result_item = {
        **item,
        "extraction_result": result,
        "extraction_method": method,
        "blocked_reason": error,
        "extracted_character_count": len(normalized),
        "extracted_line_count": normalized.count("\n") + (1 if normalized else 0),
        "text_sha256": _hash_text(normalized) if normalized else None,
        "post_extraction_classification": _post_extraction_classification(item),
        **(
            {
                "ocr_engine_available": shutil.which("tesseract") is not None,
                "ocr_requested_languages": ["chi_sim", "eng"],
                "ocr_available_languages": _ocr_available_languages(),
                "ocr_backend": "tesseract" if method == "pytesseract" else method,
                "ocr_cache_status": "miss",
            }
            if status == "image_ocr_queued"
            else {}
        ),
    }
    if status == "image_ocr_queued":
        _write_ocr_cached_item(result_item)
    return result_item


def build_extraction_results(*, write: bool = True) -> dict[str, Any]:
    queue_report = build_extraction_queue(write=False)
    results = [_extract_queued_item(item) for item in queue_report["queue"]]
    result_counts: dict[str, int] = {}
    method_counts: dict[str, int] = {}
    classification_counts: dict[str, int] = {}
    for item in results:
        result_counts[item["extraction_result"]] = result_counts.get(item["extraction_result"], 0) + 1
        method_counts[item["extraction_method"]] = method_counts.get(item["extraction_method"], 0) + 1
        classification_counts[item["post_extraction_classification"]] = (
            classification_counts.get(item["post_extraction_classification"], 0) + 1
        )
    stored_text_payload_fields = sum(1 for item in results if "text" in item or "text_preview" in item)
    report = {
        "scope": "extraction-results",
        "status": "pass",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "mode": {
            "heavy_ocr": shutil.which("tesseract") is not None,
            "macos_vision_available": _vision_ocr_binary() is not None,
            "ocr_backend_policy": "prefer_tesseract_else_macos_vision_else_blocked",
            "whole_machine_scan": False,
            "external_high_relevance_scan": True,
            "semantic_promotion": False,
            "boundary": "Extraction results store text hashes and counts only. Extracted private text is not copied into the repo truth chain.",
        },
        "summary": {
            "total_files": len(results),
            "unhashed_files": sum(1 for item in results if not item.get("sha256")),
            "stored_text_payload_fields": stored_text_payload_fields,
        },
        "result_counts": dict(sorted(result_counts.items())),
        "method_counts": dict(sorted(method_counts.items())),
        "post_extraction_classification_counts": dict(sorted(classification_counts.items())),
        "results": results,
        "artifacts": {
            "json_report": _relative(EXTRACTION_RESULTS_JSON),
            "markdown_report": _relative(EXTRACTION_RESULTS_MD),
        },
        "title": "Extraction Results Manifest",
    }
    if write:
        _write_reports(report, json_report=EXTRACTION_RESULTS_JSON, markdown_report=EXTRACTION_RESULTS_MD)
    return report


def build_extraction_queue(*, write: bool = True) -> dict[str, Any]:
    project = build_manifest(scope="project", write=False)
    external = build_manifest(scope="external", write=False)
    queue: list[dict[str, Any]] = []
    for source_scope, manifest in [("project", project), ("external", external)]:
        for item in manifest["by_path"].values():
            if item["extraction_status"] not in QUEUED_EXTRACTION_STATUSES:
                continue
            queue.append(
                {
                    "source_scope": source_scope,
                    "path": item["path"],
                    "suffix": item["suffix"],
                    "byte_count": item["byte_count"],
                    "sha256": item["sha256"],
                    "extraction_status": item["extraction_status"],
                    "classification": item["classification"],
                    "priority": item["priority"],
                    "promotion_status": item["promotion_status"],
                }
            )
    queue.sort(key=lambda item: (item["extraction_status"], item["source_scope"], item["path"]))
    queue_counts: dict[str, int] = {}
    source_counts: dict[str, int] = {}
    for item in queue:
        queue_counts[item["extraction_status"]] = queue_counts.get(item["extraction_status"], 0) + 1
        source_counts[item["source_scope"]] = source_counts.get(item["source_scope"], 0) + 1
    report = {
        "scope": "extraction-queue",
        "status": "pass",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "mode": {
            "heavy_ocr": False,
            "whole_machine_scan": False,
            "external_high_relevance_scan": True,
            "semantic_promotion": False,
            "boundary": "Queue only. It identifies PDF/image/document extraction work without performing OCR or promoting extracted text.",
        },
        "summary": {
            "queued_files": len(queue),
            "unhashed_files": sum(1 for item in queue if not item.get("sha256")),
            "source_counts": dict(sorted(source_counts.items())),
        },
        "queue_counts": dict(sorted(queue_counts.items())),
        "queue": queue,
        "artifacts": {
            "json_report": _relative(EXTRACTION_QUEUE_JSON),
            "markdown_report": _relative(EXTRACTION_QUEUE_MD),
        },
        "title": "Extraction Queue Manifest",
    }
    if write:
        _write_reports(report, json_report=EXTRACTION_QUEUE_JSON, markdown_report=EXTRACTION_QUEUE_MD)
    return report


def _write_reports(report: dict[str, Any], *, json_report: Path, markdown_report: Path) -> None:
    json_report.parent.mkdir(parents=True, exist_ok=True)
    json_report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    markdown_report.write_text(_render_markdown(report), encoding="utf-8")


def _render_markdown(report: dict[str, Any]) -> str:
    summary = report["summary"]
    if report["scope"] == "extraction-queue":
        return _render_extraction_queue_markdown(report)
    if report["scope"] == "extraction-results":
        return _render_extraction_results_markdown(report)
    lines = [
        f"# {report.get('title', 'Character-Level Inventory Manifest')}",
        "",
        f"- status: {report['status']}",
        f"- scope: {report['scope']}",
        f"- generated_at: {report['generated_at']}",
        f"- total_files: {summary['total_files']}",
        f"- unhashed_files: {summary['unhashed_files']}",
        f"- unclassified_files: {summary['unclassified_files']}",
        f"- unknown_extraction_status: {summary['unknown_extraction_status']}",
        f"- Heavy OCR: {'enabled' if report['mode']['heavy_ocr'] else 'disabled'}",
        f"- Whole-machine scan: {'enabled' if report['mode']['whole_machine_scan'] else 'disabled'}",
        "",
        "## Root Summary",
        "",
        "| Root | Files | Bytes |",
        "| --- | ---: | ---: |",
    ]
    for root, item in report["root_summary"].items():
        lines.append(f"| `{root}` | {item['files']} | {item['bytes']} |")
    lines.extend(["", "## Extraction Status Counts", ""])
    for name, count in summary["extraction_status_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(["", "## Classification Counts", ""])
    for name, count in summary["classification_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            report["mode"]["boundary"],
            "",
            "This manifest proves indexing, hashing, and extraction-state classification. It does not by itself promote any source into the runtime truth chain.",
            "",
        ]
    )
    return "\n".join(lines)


def _render_extraction_results_markdown(report: dict[str, Any]) -> str:
    lines = [
        f"# {report.get('title', 'Extraction Results Manifest')}",
        "",
        f"- status: {report['status']}",
        f"- scope: {report['scope']}",
        f"- generated_at: {report['generated_at']}",
        f"- total_files: {report['summary']['total_files']}",
        f"- unhashed_files: {report['summary']['unhashed_files']}",
        f"- stored_text_payload_fields: {report['summary']['stored_text_payload_fields']}",
        f"- Heavy OCR: {'enabled' if report['mode']['heavy_ocr'] else 'disabled'}",
        f"- macos_vision_available: {report['mode'].get('macos_vision_available', False)}",
        f"- ocr_backend_policy: {report['mode'].get('ocr_backend_policy', 'unknown')}",
        "",
        "## Result Counts",
        "",
    ]
    for name, count in report["result_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(["", "## Method Counts", ""])
    for name, count in report["method_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(["", "## Post-Extraction Classification Counts", ""])
    for name, count in report["post_extraction_classification_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            report["mode"]["boundary"],
            "",
            "No extracted text or OCR text is stored in this report; only hashes, counts, methods, and statuses are persisted.",
            "",
        ]
    )
    return "\n".join(lines)


def _render_extraction_queue_markdown(report: dict[str, Any]) -> str:
    lines = [
        f"# {report.get('title', 'Extraction Queue Manifest')}",
        "",
        f"- status: {report['status']}",
        f"- scope: {report['scope']}",
        f"- generated_at: {report['generated_at']}",
        f"- queued_files: {report['summary']['queued_files']}",
        f"- unhashed_files: {report['summary']['unhashed_files']}",
        f"- Heavy OCR: {'enabled' if report['mode']['heavy_ocr'] else 'disabled'}",
        f"- Whole-machine scan: {'enabled' if report['mode']['whole_machine_scan'] else 'disabled'}",
        "",
        "## Queue Counts",
        "",
    ]
    for name, count in report["queue_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(["", "## Source Counts", ""])
    for name, count in report["summary"]["source_counts"].items():
        lines.append(f"- `{name}`: {count}")
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            report["mode"]["boundary"],
            "",
            "Queued files are indexed by path, size, hash, and extraction status only. Private source text is not copied into this report.",
            "",
        ]
    )
    return "\n".join(lines)


def _summary_view(report: dict[str, Any]) -> dict[str, Any]:
    return {
        key: value
        for key, value in report.items()
        if key not in {"by_path"}
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--scope",
        default="project",
        choices=["project", "external", "extraction-queue", "extraction-results"],
    )
    parser.add_argument("--no-write", action="store_true", help="Print the manifest without writing report artifacts.")
    parser.add_argument("--summary-only", action="store_true", help="Print only summary fields; still writes full artifacts unless --no-write is set.")
    args = parser.parse_args(argv)

    report = build_manifest(scope=args.scope, write=not args.no_write)
    printable = _summary_view(report) if args.summary_only else report
    print(json.dumps(printable, ensure_ascii=False, indent=2))
    return 0 if report["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
